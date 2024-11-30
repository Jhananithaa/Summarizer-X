import os
import PyPDF2
import docx
import streamlit as st
from groq import Groq
from io import BytesIO
from gtts import gTTS
import pyperclip
from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet

# PDF Generation Function
def generate_pdf(content: str) -> BytesIO:
    """
    Generates a PDF with properly formatted and wrapped text.
    """
    pdf_file = BytesIO()
    doc = SimpleDocTemplate(
        pdf_file,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
    )

    styles = getSampleStyleSheet()
    story = []

    # Split the content into paragraphs and add them to the story
    for paragraph_text in content.splitlines():
        if paragraph_text.strip():  # Add only non-empty lines
            paragraph = Paragraph(paragraph_text, styles["Normal"])
            story.append(paragraph)

    # Build the PDF document
    doc.build(story)
    pdf_file.seek(0)  # Reset file pointer to the beginning
    return pdf_file

# Streamlit page configuration
st.set_page_config(
    page_title="SummarizerX - ChatBot",
    layout="centered"
)

working_dir = os.path.dirname(os.path.abspath(__file__))

GROQ_API_KEY = "ADD GROQ API KEY HERE"
os.environ["GROQ_API_KEY"] = GROQ_API_KEY
client = Groq()

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "document_content" not in st.session_state:
    st.session_state.document_content = ""

st.title("SummarizerX - ChatBot")
st.text("Chat with Our AI Chatbot.\nAsk questions, get summaries, and learn more about the content!")

# Chat history display
for idx, message in enumerate(st.session_state.chat_history):
    with st.chat_message(message["role"]):
        st.markdown(message["content"])
        if message["role"] == "assistant":
            col1, col2, col3 = st.columns([1, 1, 1])

            # Option buttons
            with col1:
                read_aloud_triggered = st.button("Read Aloud", key=f"read_aloud_{idx}")
            with col2:
                if st.button("Copy to Clipboard", key=f"copy_clipboard_{idx}"):
                    try:
                        pyperclip.copy(message["content"])
                        st.success("Response copied to clipboard!")
                    except Exception as copy_error:
                        st.error(f"Error copying to clipboard: {str(copy_error)}")
            with col3:
                if st.button("Download Options", key=f"download_options_{idx}"):
                    # TXT Download
                    st.download_button(
                        label="Download as TXT",
                        data=message["content"].encode('utf-8'),
                        file_name="assistant_response.txt",
                        mime="text/plain"
                    )
                    # DOCX Download
                    doc = docx.Document()
                    doc.add_paragraph(message["content"])
                    docx_file = BytesIO()
                    doc.save(docx_file)
                    docx_file.seek(0)
                    st.download_button(
                        label="Download as Word",
                        data=docx_file,
                        file_name="assistant_response.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    # PDF Download
                    pdf_file = generate_pdf(message["content"])
                    st.download_button(
                        label="Download as PDF",
                        data=pdf_file,
                        file_name="assistant_response.pdf",
                        mime="application/pdf"
                    )

            # Audio player (placed below the buttons)
            if read_aloud_triggered:
                try:
                    tts = gTTS(message["content"], lang='en')
                    audio_file = BytesIO()
                    tts.write_to_fp(audio_file)
                    audio_file.seek(0)
                    st.audio(audio_file, format='audio/mp3', start_time=0)
                except Exception as tts_error:
                    st.error(f"Error generating audio: {str(tts_error)}")


user_prompt = st.chat_input("Ask LLAMA...")

uploaded_file = st.file_uploader("Upload a document (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

if uploaded_file:
    if uploaded_file.type == "application/pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        document_content = ""
        for page in reader.pages:
            document_content += page.extract_text()
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        document_content = ""
        for para in doc.paragraphs:
            document_content += para.text + "\n"
    elif uploaded_file.type == "text/plain":
        document_content = uploaded_file.read().decode("utf-8")
    st.session_state.document_content = document_content
    st.success("Document uploaded and content extracted successfully.")
# Add CTAs for summarization, paraphrasing, and quiz generation
if st.session_state.document_content:
    st.subheader("Choose an Action:")

    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        summarize_triggered = st.button("Summarize", key="summarize")
    with col2:
        paraphrase_triggered = st.button("Paraphrase", key="paraphrase")
    with col3:
        quiz_triggered = st.button("Generate Quiz", key="generate_quiz")
    
    # Handle Summarizer Button
    if summarize_triggered:
        # Append the user's choice to chat history
        user_input = "Summarize the uploaded content."
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        
        # Generate chatbot response
        try:
            messages = [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "system", "content": f"Here is the document content: {st.session_state.document_content}"},
                *st.session_state.chat_history
            ]
            
            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages
            )
            assistant_response = response.choices[0].message.content
            
            # Append assistant's response to chat history
            st.session_state.chat_history.append({"role": "assistant", "content": assistant_response})
            
            # Display the response
            with st.chat_message("assistant"):
                st.markdown(assistant_response)

        except Exception as e:
            st.error(f"Error generating summary: {str(e)}")
    
    # Handle Paraphrase Button
    if paraphrase_triggered:
        user_input = "Paraphrase the uploaded content."
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        try:
            # Similar logic as above for paraphrasing
            messages = [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "system", "content": f"Here is the document content: {st.session_state.document_content}"},
                *st.session_state.chat_history
            ]
            
            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages
            )
            assistant_response = response.choices[0].message.content
            
            st.session_state.chat_history.append({"role": "assistant", "content": assistant_response})
            
            with st.chat_message("assistant"):
                st.markdown(assistant_response)

        except Exception as e:
            st.error(f"Error generating paraphrase: {str(e)}")
    
    # Handle Quiz Generator Button
    if quiz_triggered:
        user_input = "Generate quiz questions from the uploaded content."
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        try:
            # Similar logic as above for quiz generation
            messages = [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "system", "content": f"Here is the document content: {st.session_state.document_content}"},
                *st.session_state.chat_history
            ]
            
            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages
            )
            assistant_response = response.choices[0].message.content
            
            st.session_state.chat_history.append({"role": "assistant", "content": assistant_response})
            
            with st.chat_message("assistant"):
                st.markdown(assistant_response)

        except Exception as e:
            st.error(f"Error generating quiz: {str(e)}")

# Handle user queries
if user_prompt:
    st.chat_message("user").markdown(user_prompt)
    st.session_state.chat_history.append({"role": "user", "content": user_prompt})

    if st.session_state.document_content:
        messages = [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "system", "content": f"Here is the document content: {st.session_state.document_content}"},
            *st.session_state.chat_history
        ]
    else:
        messages = [
            {"role": "system", "content": "You are a helpful assistant providing general answers."},
            *st.session_state.chat_history
        ]

    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=messages
        )

        assistant_response = response.choices[0].message.content
        st.session_state.chat_history.append({"role": "assistant", "content": assistant_response})

        with st.chat_message("assistant"):
            st.markdown(assistant_response)

    except Exception as e:
        st.error(f"Error processing the question: {str(e)}")
